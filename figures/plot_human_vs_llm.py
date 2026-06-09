"""LLM-vs-human rater plot for the last 10 papers in scoring_guide_ana_v2.docx.

Extracts Ana's hand-scored B/E/D item scores from the docx, pairs each
human-scored assay to the matching LLM consensus assay (by fuzzy name),
and renders:
  Row 1: assay-level scatter (human vs LLM) for B, E, D totals.
  Row 2: per-item pooled confusion heatmaps for B, E, D.
Plus per-item weighted Cohen's kappa and ICC printed to stdout.
"""
from __future__ import annotations

import json
import re
from collections import defaultdict
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

import matplotlib.pyplot as plt
import numpy as np

_DATA = Path(__file__).resolve().parent.parent / "data"
_OUT = Path(__file__).resolve().parent / "output"
DOCX = _DATA / "scoring_guide_ana_v2.docx"
CONSENSUS_DIR = _DATA / "results_v2_full_consensus"
OUT_PNG = _OUT / "human_vs_llm.png"
OUT_PER_ASSAY = _OUT / "human_vs_llm_per_assay.png"

# Item → score column index mapping in the consensus JSON. Names follow the
# B1_..., E1_..., D1_... prefix pattern.
B_ITEMS = ["B1", "B2", "B3", "B4", "B5", "B6", "B7", "B8"]
E_ITEMS = ["E1", "E2", "E3", "E4", "E5", "E6", "E7", "E8"]
D_ITEMS = ["D1", "D2", "D3", "D4", "D5", "D6"]

PAPERS_BOTTOM_10 = [
    "horrocks2025", "huang2022", "jeanblanc2024", "kato2025", "lee2020",
    "marcherrorsted2020", "marek2018", "popik2019", "wang2025", "yu2023",
]

PAPER_RE = re.compile(r"^Paper:\s*(\d+)\.\s*(\S+)")


def parse_docx(path: Path) -> list[dict]:
    """Walk the docx body in order, grouping paragraphs+tables under each paper."""
    d = Document(str(path))
    papers, current = [], None
    para_iter = iter(d.paragraphs)
    tbl_iter = iter(d.tables)
    for child in d.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = next(para_iter)
            txt = p.text.strip()
            m = PAPER_RE.match(txt)
            if m:
                if current:
                    papers.append(current)
                current = {
                    "num": int(m.group(1)),
                    "stem": m.group(2),
                    "paras": [],
                    "tables": [],
                }
            elif current is not None:
                current["paras"].append(txt)
        elif child.tag == qn("w:tbl"):
            t = next(tbl_iter)
            if current is not None:
                current["tables"].append(t)
    if current:
        papers.append(current)
    return papers


def _score_cell(row) -> str:
    """Return the score-write cell text (col 2 for score sheets, col 2 for housing/exp)."""
    return row.cells[2].text.strip()


def _row_item_label(row) -> str:
    return row.cells[0].text.strip().split(".")[0].strip()


def extract_paper(paper: dict) -> dict:
    """Pull housing, and per-assay exp_cond + B/E/D items from the table sequence."""
    tables = paper["tables"]
    paras = paper["paras"]
    if not tables:
        return None
    # Tables: [housing] [exp_cond, B, E, D]*n_assays
    housing_tbl = tables[0]
    assay_tbls = tables[1:]
    if len(assay_tbls) % 4 != 0:
        print(f"WARN {paper['stem']}: assay tables not multiple of 4: {len(assay_tbls)}")

    # Housing — H1..H4
    housing = {}
    for r in housing_tbl.rows[1:]:
        label = _row_item_label(r)
        val = _score_cell(r)
        if val:
            housing[label] = val.lower()

    # Pull assay names from paragraphs ("Assay name: ___ NAME ___")
    name_re = re.compile(r"Assay name:\s*[_\s]*(.+?)[_\s]*$")
    assay_names = []
    for p in paras:
        m = name_re.match(p)
        if m:
            nm = m.group(1).strip("_ ").strip()
            assay_names.append(nm if nm else None)

    assays = []
    n_assays = len(assay_tbls) // 4
    for i in range(n_assays):
        exp_t, b_t, e_t, d_t = assay_tbls[4 * i : 4 * i + 4]
        exp_cond = {}
        for r in exp_t.rows[1:]:
            v = _score_cell(r)
            if v:
                exp_cond[_row_item_label(r)] = v.lower()
        b_scores, e_scores, d_scores = {}, {}, {}
        for r in b_t.rows[1:]:
            v = _score_cell(r)
            if v != "":
                b_scores[_row_item_label(r)] = int(v) if v.isdigit() else None
        for r in e_t.rows[1:]:
            v = _score_cell(r)
            if v != "":
                e_scores[_row_item_label(r)] = int(v) if v.isdigit() else None
        for r in d_t.rows[1:]:
            v = _score_cell(r)
            if v != "":
                d_scores[_row_item_label(r)] = int(v) if v.isdigit() else None
        name = assay_names[i] if i < len(assay_names) else None
        # Skip totally blank assay sheets (placeholders)
        if not (b_scores or e_scores or d_scores or exp_cond or name):
            continue
        assays.append({
            "name": name,
            "exp_cond": exp_cond,
            "B": b_scores,
            "E": e_scores,
            "D": d_scores,
        })

    return {"stem": paper["stem"], "num": paper["num"], "housing": housing, "assays": assays}


def llm_items(assay: dict, prefix: str) -> dict:
    """Pull {B1, B2, ...: score} from a consensus assay block."""
    sect = assay.get({"B": "behavioural_complexity", "E": "environmental_complexity", "D": "recording_duration"}[prefix], {})
    out = {}
    for k, v in sect.items():
        m = re.match(rf"({prefix}\d+)_", k)
        if m and isinstance(v, dict) and "score" in v:
            out[m.group(1)] = int(v["score"])
    return out


ASSAY_ALIASES = {
    "oft": "open field test locomotor activity",
    "epm": "elevated plus maze",
    "fst": "forced swim test",
    "spt": "sucrose preference test",
    "nor": "novel object recognition",
    "nort": "novel object recognition",
    "ppi": "prepulse inhibition",
    "htr": "head twitch response",
    "wds": "head twitch response",
    "wet-dog shake": "head twitch response",
    "cpa": "conditioned place aversion",
    "cpp": "conditioned place preference",
    "mbt": "marble burying test",
    "tst": "tail suspension test",
}


def _normalize_name(s: str) -> str:
    s = s.lower()
    # expand standalone abbreviation tokens
    tokens = re.split(r"[\s_/\-]+", s)
    expanded = [ASSAY_ALIASES.get(t, t) for t in tokens]
    out = " ".join(expanded)
    # also expand inline phrases
    for k, v in ASSAY_ALIASES.items():
        if k in out and len(k) > 2:
            out = out.replace(k, v)
    return out


def _pair_score(human_name: str, llm_name: str) -> float:
    h = _normalize_name(human_name)
    l = _normalize_name(llm_name)
    ht = set(re.findall(r"\w+", h))
    lt = set(re.findall(r"\w+", l))
    overlap = len(ht & lt) / max(1, len(ht | lt))
    ratio = SequenceMatcher(None, h, l).ratio()
    return 0.5 * overlap + 0.5 * ratio


def match_assays(human_assays: list[dict], llm_assays: list[dict], threshold: float = 0.4) -> list[tuple[int, int]]:
    """Greedy best-first matching. Returns list of (human_idx, llm_idx) pairs."""
    candidates = []
    for hi, ha in enumerate(human_assays):
        if not ha.get("name"):
            continue
        for li, la in enumerate(llm_assays):
            ln = la.get("assay_name") or ""
            sc = _pair_score(ha["name"], ln)
            if sc >= threshold:
                candidates.append((sc, hi, li))
    candidates.sort(reverse=True)
    used_h, used_l, pairs = set(), set(), []
    for sc, hi, li in candidates:
        if hi in used_h or li in used_l:
            continue
        used_h.add(hi); used_l.add(li); pairs.append((hi, li))
    return pairs


def total_b(items: dict) -> int | None:
    if not items or any(items.get(k) is None for k in B_ITEMS):
        return None
    return sum(items[k] for k in B_ITEMS)


def total_e(items: dict) -> int | None:
    if not items or any(items.get(k) is None for k in E_ITEMS):
        return None
    return sum(items[k] for k in E_ITEMS)


def total_d(items: dict) -> int | None:
    if not items or any(items.get(k) is None for k in D_ITEMS):
        return None
    return sum(items[k] for k in D_ITEMS)


def weighted_kappa(human: list[int], llm: list[int], maxv: int, weights: str = "linear") -> float | None:
    """Cohen's weighted kappa for ordinal scores on 0..maxv."""
    if not human or len(human) != len(llm):
        return None
    n = len(human)
    cats = list(range(maxv + 1))
    k = len(cats)
    O = np.zeros((k, k))
    for a, b in zip(human, llm):
        if a is None or b is None:
            continue
        O[a, b] += 1
    n_eff = O.sum()
    if n_eff == 0:
        return None
    row_marg = O.sum(axis=1, keepdims=True)
    col_marg = O.sum(axis=0, keepdims=True)
    E = row_marg @ col_marg / n_eff
    W = np.zeros((k, k))
    for i in range(k):
        for j in range(k):
            if weights == "linear":
                W[i, j] = abs(i - j) / (k - 1)
            else:  # quadratic
                W[i, j] = ((i - j) / (k - 1)) ** 2
    num = (W * O).sum()
    den = (W * E).sum()
    if den == 0:
        return 1.0 if num == 0 else None
    return 1 - num / den


def icc_2_1(rater1: list[int], rater2: list[int]) -> float | None:
    """ICC(2,1) two-way random, single measure, absolute agreement."""
    if not rater1 or len(rater1) != len(rater2):
        return None
    Y = np.array([[a, b] for a, b in zip(rater1, rater2)], dtype=float)
    n, k = Y.shape  # n subjects, k raters (=2)
    if n < 2:
        return None
    mean_subj = Y.mean(axis=1)
    mean_rater = Y.mean(axis=0)
    grand = Y.mean()
    MSR = k * np.sum((mean_subj - grand) ** 2) / (n - 1)
    MSC = n * np.sum((mean_rater - grand) ** 2) / (k - 1)
    SSE = np.sum((Y - mean_subj[:, None] - mean_rater[None, :] + grand) ** 2)
    MSE = SSE / ((n - 1) * (k - 1))
    denom = MSR + (k - 1) * MSE + k * (MSC - MSE) / n
    if denom <= 0:
        return None
    return (MSR - MSE) / denom


def spearman(x: list[float], y: list[float]) -> float | None:
    if not x or len(x) < 2:
        return None
    def rank(a):
        order = np.argsort(a, kind="mergesort")
        r = np.empty_like(order, dtype=float)
        r[order] = np.arange(1, len(a) + 1)
        # average ties
        a_arr = np.asarray(a)
        for v in set(a):
            idx = np.where(a_arr == v)[0]
            if len(idx) > 1:
                r[idx] = r[idx].mean()
        return r
    rx, ry = rank(x), rank(y)
    return float(np.corrcoef(rx, ry)[0, 1])


def main():
    docx_papers = parse_docx(DOCX)
    # Restrict to bottom 10
    docx_papers = [p for p in docx_papers if p["stem"] in PAPERS_BOTTOM_10]
    print(f"Parsed {len(docx_papers)} human-scored papers from docx.")

    pairs = []  # one entry per matched assay
    for paper in docx_papers:
        human = extract_paper(paper)
        if human is None or not human["assays"]:
            continue
        stem = human["stem"]
        with open(CONSENSUS_DIR / f"{stem}.json") as f:
            consensus = json.load(f)
        llm_assays = consensus.get("assays", [])
        matched = match_assays(human["assays"], llm_assays)
        matched_h = {hi for hi, _ in matched}
        for hi, ha in enumerate(human["assays"]):
            if hi not in matched_h:
                print(f"  skip {stem} / {ha['name']!r}: no LLM match")
        for hi, li in matched:
            ha = human["assays"][hi]
            llm_assay = llm_assays[li]
            pairs.append({
                "stem": stem,
                "human_name": ha["name"],
                "llm_name": llm_assay.get("assay_name"),
                "human_B": ha["B"],
                "human_E": ha["E"],
                "human_D": ha["D"],
                "llm_B": llm_items(llm_assay, "B"),
                "llm_E": llm_items(llm_assay, "E"),
                "llm_D": llm_items(llm_assay, "D"),
            })

    print(f"Matched {len(pairs)} assay pairs.\n")
    for p in pairs:
        hb, lb = total_b(p["human_B"]), total_b(p["llm_B"])
        he, le = total_e(p["human_E"]), total_e(p["llm_E"])
        hd, ld = total_d(p["human_D"]), total_d(p["llm_D"])
        print(f"  {p['stem']:22} h='{(p['human_name'] or '?')[:35]:35}' l='{(p['llm_name'] or '?')[:35]:35}' "
              f"B {hb}/{lb}  E {he}/{le}  D {hd}/{ld}")

    # ------ Compute stats ------
    dims = {
        "B": (B_ITEMS, "human_B", "llm_B"),
        "E": (E_ITEMS, "human_E", "llm_E"),
        "D": (D_ITEMS, "human_D", "llm_D"),
    }
    item_max = {**{k: 3 for k in ("B1",)}, **{k: 1 for k in ("B2","B3","B4","B6","B7","B8")}, "B5": 3,
                "E1": 3, **{k: 1 for k in ("E2","E3","E4","E5","E6","E8")}, "E7": 2,
                "D1": 2, "D2": 3, "D3": 3, **{k: 1 for k in ("D4","D5","D6")}}

    print("\n=== Per-dimension agreement (assay-level raw totals) ===")
    totals = {}
    for dim, (items, hk, lk) in dims.items():
        H, L = [], []
        for p in pairs:
            h = sum(p[hk].get(it) for it in items) if all(p[hk].get(it) is not None for it in items) else None
            l = sum(p[lk].get(it) for it in items) if all(p[lk].get(it) is not None for it in items) else None
            if h is not None and l is not None:
                H.append(h); L.append(l)
        totals[dim] = (H, L)
        n = len(H)
        exact = sum(1 for a, b in zip(H, L) if a == b)
        diff = np.array(H) - np.array(L)
        rho = spearman(H, L)
        icc = icc_2_1(H, L)
        print(f"  {dim}: n={n}, exact={exact}/{n} ({100*exact/n:.0f}%), "
              f"mean diff (H−L)={diff.mean():+.2f} (sd {diff.std(ddof=1):.2f}), "
              f"Spearman ρ={rho:.2f}, ICC(2,1)={icc:.2f}")

    print("\n=== Per-item linearly-weighted Cohen's kappa ===")
    item_stats = {}
    for dim, (items, hk, lk) in dims.items():
        for it in items:
            H = [p[hk].get(it) for p in pairs if p[hk].get(it) is not None and p[lk].get(it) is not None]
            L = [p[lk].get(it) for p in pairs if p[hk].get(it) is not None and p[lk].get(it) is not None]
            kap = weighted_kappa(H, L, item_max[it])
            exact = sum(1 for a, b in zip(H, L) if a == b)
            n = len(H)
            item_stats[it] = (n, exact, kap)
            print(f"  {it}: n={n}, exact={exact}/{n} ({100*exact/n:.0f}%), κ_linear={kap:.2f}" if kap is not None
                  else f"  {it}: n={n}, exact={exact}/{n} ({100*exact/n:.0f}%), κ=NA")

    # ------ Plot ------
    fig, axes = plt.subplots(2, 3, figsize=(13.5, 8.5))
    dim_colour = {"B": "#7B2FBE", "E": "#E6550D", "D": "#3a7acf"}
    dim_full = {"B": "Behavioural complexity", "E": "Environmental complexity", "D": "Recording duration"}

    rng = np.random.default_rng(20260602)
    for col, dim in enumerate(["B", "E", "D"]):
        ax = axes[0, col]
        H, L = totals[dim]
        H_arr = np.array(H, dtype=float)
        L_arr = np.array(L, dtype=float)
        jx = rng.uniform(-0.12, 0.12, size=len(H))
        jy = rng.uniform(-0.12, 0.12, size=len(L))
        ax.plot([0, 11], [0, 11], "-", color="#bbb", lw=1)
        ax.scatter(H_arr + jx, L_arr + jy, s=44, color=dim_colour[dim], alpha=0.78,
                   edgecolor="white", linewidth=0.8)
        n = len(H)
        exact = sum(1 for a, b in zip(H, L) if a == b)
        rho = spearman(H, L) or 0
        icc = icc_2_1(H, L) or 0
        ax.set_xlim(-0.5, 11.5); ax.set_ylim(-0.5, 11.5)
        ax.set_xticks(range(0, 12, 2)); ax.set_yticks(range(0, 12, 2))
        ax.set_xlabel("Human (raw total)")
        ax.set_ylabel("LLM consensus (raw total)")
        ax.set_title(f"{dim_full[dim]}", fontsize=11, fontweight="bold")
        ax.text(0.04, 0.96, f"n={n}\nexact = {exact}/{n} ({100*exact/n:.0f}%)\nρ = {rho:.2f}\nICC = {icc:.2f}",
                transform=ax.transAxes, va="top", ha="left", fontsize=9,
                bbox=dict(facecolor="white", edgecolor="#ddd", boxstyle="round,pad=0.4"))
        ax.set_aspect("equal", "box")
        ax.grid(True, color="#eee", linewidth=0.7)

        # Bottom row: pooled per-item confusion (rescale every item to its own 0..1)
        items = dims[dim][0]
        hk, lk = dims[dim][1], dims[dim][2]
        # use raw item scores 0..3 (max across items in this dim)
        max_dim_item = max(item_max[it] for it in items)
        M = np.zeros((max_dim_item + 1, max_dim_item + 1), dtype=int)
        for p in pairs:
            for it in items:
                h, l = p[hk].get(it), p[lk].get(it)
                if h is None or l is None:
                    continue
                M[h, l] += 1
        ax2 = axes[1, col]
        im = ax2.imshow(M, origin="lower", cmap="Purples" if dim == "B" else ("Oranges" if dim == "E" else "Blues"))
        for i in range(M.shape[0]):
            for j in range(M.shape[1]):
                if M[i, j] > 0:
                    ax2.text(j, i, str(M[i, j]), ha="center", va="center",
                             color="white" if M[i, j] > M.max() * 0.55 else "#333", fontsize=9)
        ax2.set_xlabel("LLM item score")
        ax2.set_ylabel("Human item score")
        ax2.set_xticks(range(max_dim_item + 1))
        ax2.set_yticks(range(max_dim_item + 1))
        ax2.set_title(f"Pooled {dim}-item confusion", fontsize=10)
        # Diagonal guide
        ax2.plot([-0.5, max_dim_item + 0.5], [-0.5, max_dim_item + 0.5], "-", color="#999", lw=0.8)
        ax2.set_xlim(-0.5, max_dim_item + 0.5); ax2.set_ylim(-0.5, max_dim_item + 0.5)

    fig.suptitle(f"Human (Ana) vs LLM consensus — {len(pairs)} assays, papers 11–20 of scoring_guide_ana_v2",
                 fontsize=12, fontweight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.96])
    fig.savefig(OUT_PNG, dpi=160)
    print(f"\nWrote {OUT_PNG}")

    plot_per_assay(pairs, item_max)


def plot_per_assay(pairs: list[dict], item_max: dict) -> None:
    """Grid of small multiples: one panel per assay, item-level human vs LLM."""
    n = len(pairs)
    ncols = 2
    nrows = (n + ncols - 1) // ncols
    fig, axes = plt.subplots(nrows, ncols, figsize=(15, 2.6 * nrows + 1.0), squeeze=False)
    dim_colour = {"B": "#7B2FBE", "E": "#E6550D", "D": "#3a7acf"}
    items_in_order = [(it, "B") for it in B_ITEMS] + [(it, "E") for it in E_ITEMS] + [(it, "D") for it in D_ITEMS]
    item_labels = [it for it, _ in items_in_order]
    x = np.arange(len(items_in_order))
    bar_w = 0.4

    for idx, p in enumerate(pairs):
        ax = axes[idx // ncols][idx % ncols]
        humans, llms, colours = [], [], []
        for it, dim in items_in_order:
            h = p[f"human_{dim}"].get(it)
            l = p[f"llm_{dim}"].get(it)
            humans.append(h if h is not None else np.nan)
            llms.append(l if l is not None else np.nan)
            colours.append(dim_colour[dim])
        humans_arr = np.array(humans, dtype=float)
        llms_arr = np.array(llms, dtype=float)
        # Background shading by dimension
        for span_dim, (lo, hi) in [("B", (0, len(B_ITEMS))), ("E", (len(B_ITEMS), len(B_ITEMS) + len(E_ITEMS))), ("D", (len(B_ITEMS) + len(E_ITEMS), len(items_in_order)))]:
            ax.axvspan(lo - 0.5, hi - 0.5, facecolor=dim_colour[span_dim], alpha=0.05)
        # Human = filled, LLM = open
        ax.bar(x - bar_w / 2, humans_arr, width=bar_w, color=colours, alpha=0.85, edgecolor="none", label="Human")
        ax.bar(x + bar_w / 2, llms_arr, width=bar_w, color="white", edgecolor=colours, linewidth=1.4, label="LLM")
        # Mark disagreements with a small red ✕ at the higher of the two
        for xi, (h_, l_) in enumerate(zip(humans_arr, llms_arr)):
            if np.isnan(h_) or np.isnan(l_):
                continue
            if h_ != l_:
                ax.scatter(xi, max(h_, l_) + 0.18, marker="x", color="#d33", s=22, linewidth=1.4, zorder=5)

        # Per-dim raw totals in title
        b_h = sum(humans[i] for i in range(len(B_ITEMS)) if not np.isnan(humans[i]))
        b_l = sum(llms[i] for i in range(len(B_ITEMS)) if not np.isnan(llms[i]))
        e_h = sum(humans[i] for i in range(len(B_ITEMS), len(B_ITEMS) + len(E_ITEMS)) if not np.isnan(humans[i]))
        e_l = sum(llms[i] for i in range(len(B_ITEMS), len(B_ITEMS) + len(E_ITEMS)) if not np.isnan(llms[i]))
        d_h = sum(humans[i] for i in range(len(B_ITEMS) + len(E_ITEMS), len(items_in_order)) if not np.isnan(humans[i]))
        d_l = sum(llms[i] for i in range(len(B_ITEMS) + len(E_ITEMS), len(items_in_order)) if not np.isnan(llms[i]))
        ax.set_title(f"{p['stem']} · {(p['human_name'] or p['llm_name'] or '?')[:40]}\n"
                     f"B {int(b_h)}/{int(b_l)} · E {int(e_h)}/{int(e_l)} · D {int(d_h)}/{int(d_l)}  (H/L)",
                     fontsize=9, loc="left")
        ax.set_xticks(x)
        ax.set_xticklabels(item_labels, rotation=0, fontsize=8)
        ax.set_xlim(-0.7, len(items_in_order) - 0.3)
        ax.set_ylim(0, 3.6)
        ax.set_yticks([0, 1, 2, 3])
        ax.grid(axis="y", color="#eee", linewidth=0.6)
        ax.set_axisbelow(True)
        if idx % ncols == 0:
            ax.set_ylabel("score")

    # Hide unused panels
    for empty in range(n, nrows * ncols):
        axes[empty // ncols][empty % ncols].axis("off")

    # Legend (single, outside)
    from matplotlib.patches import Patch
    legend_handles = [
        Patch(facecolor="#7B2FBE", edgecolor="none", label="B (filled = Human)"),
        Patch(facecolor="#E6550D", edgecolor="none", label="E"),
        Patch(facecolor="#3a7acf", edgecolor="none", label="D"),
        Patch(facecolor="white", edgecolor="#444", label="open bar = LLM"),
    ]
    fig.suptitle("Per-assay item scores — Human (filled) vs LLM (open); ✕ marks disagreement",
                 fontsize=12, fontweight="bold", y=0.995)
    fig.legend(handles=legend_handles, loc="upper center", ncol=4, frameon=False, fontsize=9, bbox_to_anchor=(0.5, 0.978))
    fig.tight_layout(rect=[0, 0, 1, 0.96])
    fig.savefig(OUT_PER_ASSAY, dpi=150)
    print(f"Wrote {OUT_PER_ASSAY}")


if __name__ == "__main__":
    main()
